"""程序说明"""
# -*-  coding: utf-8 -*-
# Author: cao wang
# Datetime : 2020
# software: PyCharm
# 收获:

import logging
import time
#from 常用设置.文字处理.文件格式转化.读取ddp import check_file_types as check
#from 常用设置.文字处理.docx读取.docx内容 import  read_docx as read
import os
import glob
import docx
import re
import pprint
import mammoth
from lxml import etree


class Reference_Find():
    def __init__(self,input):
        self.input = input



    def docx_html(self):
        """这是由于脚注无法读取"""
        """转为html"""
        with open(self.input, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value  # The generated HTML
            #messages = result.messages  # Any messages, such as warnings during conversion
            temp = self.input.split(".")[0]+".html"
            with open(temp, "a")as f:
                f.write(html)
        return temp


    def references_write_to_docx(self):
        """首先提取脚注，然后写入"""
        temp = self.docx_html()
        print(temp)
        parser = etree.HTMLParser(encoding="gbk")
        text = etree.parse(temp, parser=parser)
        content = text.xpath("/html/body/ol[10]//text()")
        #print(content)
        references = []
        for i in content:
            if i == "↑" or "说明：" in i:
                pass
            else:
                try:
                    text = i.split("；")
                    for i in text:
                        references.append(i.replace("参见",""))
                except:
                    references.append(i.replace("参见", ""))
        #写入docx
        document = docx.Document(self.input)
        docText2 = [paragraph.text for paragraph in document.paragraphs]
        p = document.paragraphs[len(docText2)-1]
        #document.add_page_break()  # 另起一页
        p.add_run("\n")
        p.add_run("参考文献\n")
        for i,text in enumerate(references):
            i = i+1
            p.add_run("[%d]"%i+text+"\n")
        document.save(self.input)


input=r"C:\Users\lenovo\Desktop\法官量刑的限度：从危险驾驶罪的量刑情节影响因子谈起 - 副本.docx"
Reference_Find(input).references_write_to_docx()



